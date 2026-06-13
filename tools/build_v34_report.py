"""Build the formatted Mini-GRP v3.4 Word report from Markdown."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Mini-GRP-v34-Complete-Project-Report.md"
OUTPUT = ROOT / "docs" / "Mini-GRP-v34-Complete-Project-Report.docx"

NAVY = RGBColor(22, 52, 78)
BLUE = RGBColor(46, 116, 181)
TEAL = RGBColor(15, 118, 110)
GRAY = RGBColor(91, 101, 113)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
FONT = "Microsoft YaHei"


def set_run_font(run, size=11, bold=None, italic=None, color=None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    tokens = {
        "Title": (30, NAVY, 0, 8),
        "Subtitle": (14, GRAY, 0, 10),
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=50, start=120, bottom=50, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, 9, color=GRAY)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, 9, color=GRAY)


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header.add_run("MINI-GRP v3.4 | 完整项目报告")
        set_run_font(header_run, 9, bold=True, color=GRAY)
        add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = kicker.add_run("QUANTITATIVE RESEARCH SYSTEM")
    set_run_font(run, 10, bold=True, color=TEAL)
    kicker.paragraph_format.space_after = Pt(18)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Mini-GRP v3.4")
    set_run_font(run, 30, bold=True, color=NAVY)

    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("完整项目报告")
    set_run_font(run, 18, bold=True, color=BLUE)

    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead.paragraph_format.space_before = Pt(8)
    lead.paragraph_format.space_after = Pt(40)
    run = lead.add_run("从面试原型到真实数据、模型治理与 Google Cloud Run 部署")
    set_run_font(run, 12, color=GRAY)

    meta = [
        ("报告版本", "v3.4.0"),
        ("报告日期", "2026-06-13"),
        ("项目定位", "可解释的多因子研究系统与受治理的机器学习实验平台"),
        ("代码仓库", "github.com/SuperMarioGotze27/mini-grp-v32"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 7160])
    for row_index, (label, value) in enumerate(meta):
        set_cell_shading(table.cell(row_index, 0), LIGHT_BLUE)
        for column, text in enumerate((label, value)):
            paragraph = table.cell(row_index, column).paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(text)
            set_run_font(run, 10.5, bold=column == 0, color=NAVY if column == 0 else None)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(26)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run("研究、教育与面试展示用途；不构成投资建议")
    set_run_font(run, 9.5, italic=True, color=GRAY)
    doc.add_page_break()


def add_contents(doc: Document, lines: list[str]) -> None:
    doc.add_heading("内容导航", level=1)
    for line in lines:
        match = re.match(r"^(#{2,3})\s+(.+)$", line)
        if not match:
            continue
        level = len(match.group(1))
        text = match.group(2)
        if level == 2:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(text)
            set_run_font(run, 10.5, bold=True, color=NAVY)
        elif level == 3:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.45)
            paragraph.paragraph_format.space_after = Pt(2)
            run = paragraph.add_run(text)
            set_run_font(run, 9.5, color=GRAY)
    doc.add_page_break()


def create_numbering_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "290")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    p_pr.append(num_pr)


def add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\[[^\]]+\]\([^)]+\))")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            set_run_font(run)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, 10, color=NAVY)
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, bold=True)
        else:
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            run = paragraph.add_run(f"{label} ({url})")
            set_run_font(run, color=BLUE)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_run_font(run)


def add_table(doc: Document, block: list[str]) -> None:
    rows = []
    for line in block:
        values = [value.strip() for value in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", value) for value in values):
            continue
        rows.append(values)
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    normalized = [row + [""] * (column_count - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=column_count)
    table.style = "Table Grid"
    if column_count == 2:
        widths = [2800, 6560]
    elif column_count == 3:
        widths = [2200, 2380, 4780]
    elif column_count == 4:
        widths = [1500, 2300, 2300, 3260]
    else:
        base = 9360 // column_count
        widths = [base] * column_count
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)
    for row_index, row in enumerate(normalized):
        for column_index, text in enumerate(row):
            cell = table.cell(row_index, column_index)
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, LIGHT_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(text)
            set_run_font(run, 9.3, bold=row_index == 0, color=NAVY if row_index == 0 else None)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def build() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)
    add_cover(doc)
    add_contents(doc, lines)

    index = 1
    in_code = False
    code_lines: list[str] = []
    current_numbering_id: int | None = None
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.right_indent = Inches(0.25)
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(8)
                paragraph.paragraph_format.keep_together = True
                set_cell_like = OxmlElement("w:shd")
                set_cell_like.set(qn("w:fill"), LIGHT_GRAY)
                paragraph._p.get_or_add_pPr().append(set_cell_like)
                run = paragraph.add_run("\n".join(code_lines))
                set_run_font(run, 9, color=NAVY)
                run.font.name = "Consolas"
                run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
                run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Consolas")
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            current_numbering_id = None
            index += 1
            continue
        if stripped.startswith("|"):
            block = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index])
                index += 1
            add_table(doc, block)
            current_numbering_id = None
            continue
        if stripped.startswith("> "):
            current_numbering_id = None
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(10)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "F4F6F9")
            paragraph._p.get_or_add_pPr().append(shading)
            quote_text = re.sub(
                r"\[([^\]]+)\]\(([^)]+)\)",
                lambda match: f"{match.group(1)} ({match.group(2)})",
                stripped[2:],
            )
            run = paragraph.add_run(quote_text)
            set_run_font(run, 10.5, italic=True, color=NAVY)
        elif stripped.startswith("### "):
            current_numbering_id = None
            heading = doc.add_heading(stripped[4:], level=2)
            if stripped in {"### 5.2 政策权重", "### 7.4 线性真实快照回测"}:
                heading.paragraph_format.page_break_before = True
        elif stripped.startswith("## "):
            current_numbering_id = None
            heading = doc.add_heading(stripped[3:], level=1)
        elif re.match(r"^\d+\.\s+", stripped):
            if current_numbering_id is None:
                current_numbering_id = create_numbering_id(doc)
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, current_numbering_id)
            add_inline_markdown(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith("- "):
            current_numbering_id = None
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, stripped[2:])
        else:
            current_numbering_id = None
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline_markdown(paragraph, stripped)
        index += 1

    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
    doc.core_properties.title = "Mini-GRP v3.4 完整项目报告"
    doc.core_properties.subject = "多因子研究、机器学习治理与 Google Cloud Run 部署"
    doc.core_properties.author = "Mini-GRP Project"
    doc.core_properties.keywords = "Mini-GRP, quantitative research, factor model, machine learning, Cloud Run"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
